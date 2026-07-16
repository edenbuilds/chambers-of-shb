"""Post-pandoc fix script — force column widths on every table in a rendered .docx.

Pandoc pipe-tables ignore the reference.docx's tblLayout=fixed lock for column
widths and instead apportion widths based on the column-header text length —
which produces narrow stacking columns (the v0.2.0 Index-table defect).

This script opens the rendered .docx, walks every table, and applies the
column-width profile that matches the table's column count:

  5 columns  → Sr.No | Annx | Particulars   | Date | Pgs    → 8% / 8% / 60% / 14% / 10%
  4 columns  → Sr.No |      | Particulars   | Date         → 10% / 10% / 65% / 15%
  3 columns  → Sr.No | Particulars | Date                  → 10% / 75% / 15%
  2 columns  → Dates | Events                              → 18% / 82%

Also locks first-row bold + centered + vertically-centered cells across all
tables. Run as the final step of the Drafter's pandoc conversion.

Usage:
    python3 fix_docx_tables.py <input.docx> [<output.docx>]

If <output.docx> is omitted, overwrites <input.docx> in place.
"""
import sys
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# A4 usable width = 21cm - 4cm left - 2.5cm right = 14.5cm
USABLE_WIDTH_CM = 14.5

# Column-width profiles (proportions sum to 100)
PROFILES = {
    5: [8, 8, 60, 14, 10],         # Sr.No | Annx | Particulars | Date | Pgs
    4: [10, 10, 65, 15],           # Sr.No | Annx | Particulars | Date
    3: [10, 75, 15],               # Sr.No | Particulars | Date
    2: [18, 82],                   # Dates | Events
    1: [100],
}


def lock_table_layout_fixed(table):
    """Set tblLayout=fixed so column widths are honoured."""
    tblPr = table._tbl.tblPr
    # Remove any existing tblLayout
    existing = tblPr.findall(qn("w:tblLayout"))
    for el in existing:
        tblPr.remove(el)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)


def set_column_widths(table, widths_cm):
    """Apply column widths in cm to every cell in the table + tblGrid."""
    # Update tblGrid (Word's column-width table)
    tblGrid = table._tbl.find(qn("w:tblGrid"))
    if tblGrid is not None:
        # Remove existing gridCol elements
        for grid_col in tblGrid.findall(qn("w:gridCol")):
            tblGrid.remove(grid_col)
        # Add new gridCol elements with computed widths in twips (1cm = 567 twips)
        for width_cm in widths_cm:
            grid_col = OxmlElement("w:gridCol")
            grid_col.set(qn("w:w"), str(int(width_cm * 567)))
            tblGrid.append(grid_col)

    # Update every cell in every row
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_cm):
                cell.width = Cm(widths_cm[idx])


def lock_header_row(table):
    """Make the first row bold + centered + vertically-centered."""
    if not table.rows:
        return
    hdr_row = table.rows[0]
    for cell in hdr_row.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            for run in paragraph.runs:
                run.bold = True
                if not run.font.name:
                    run.font.name = "Times New Roman"
                if not run.font.size:
                    run.font.size = Pt(14)


def lock_all_cells(table):
    """Apply consistent cell formatting (no auto-indent, vertically centered)."""
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Cm(0)


def fix_table(table):
    n_cols = len(table.columns)
    if n_cols not in PROFILES:
        # Unknown column count — apply uniform widths
        widths_pct = [100 / n_cols] * n_cols
    else:
        widths_pct = PROFILES[n_cols]
    widths_cm = [USABLE_WIDTH_CM * p / 100 for p in widths_pct]

    lock_table_layout_fixed(table)
    set_column_widths(table, widths_cm)
    lock_all_cells(table)
    lock_header_row(table)


def main():
    if len(sys.argv) < 2:
        print("Usage: fix_docx_tables.py <input.docx> [<output.docx>]")
        sys.exit(1)
    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2]) if len(sys.argv) >= 3 else input_path

    doc = Document(str(input_path))
    table_count = len(doc.tables)
    for table in doc.tables:
        fix_table(table)
    doc.save(str(output_path))
    print(f"OK · fixed {table_count} table(s) in {output_path}")


if __name__ == "__main__":
    main()
