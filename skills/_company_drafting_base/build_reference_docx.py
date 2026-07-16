"""Build a filing-grade Bombay-HC reference.docx with locked Word styles.

Matches the structural conventions of a validated filed Second Appeal pleading
from the Bombay HC (your bench), supplied as the v0.2.3 filing-grade gold
standard:

- Body (Normal): TNR 14pt, 1.5 line spacing, justified, 0.5cm first-line indent
- Heading 1 (court header / case-number / cover-page anchors): TNR 14pt bold
  centered, NOT underlined (court header convention)
- Heading 2 (spaced section headers — F A C T S / G R O U N D S / P R A Y E R /
  I N D E X / S Y N O P S I S / L I S T   O F   A N N E X U R E S): TNR 14pt
  bold + UNDERLINED + centered + letter-spacing
- Heading 3 (unspaced section headers — SUBSTANTIAL QUESTIONS OF LAW /
  ACTS & RULES / CITATIONS / MOST RESPECTFULLY SHEWETH): TNR 14pt bold +
  UNDERLINED + left or centered (per Drafter choice)
- Title style: TNR 14pt bold centered (for statutory opening line like
  SECOND APPEAL UNDER SECTION 100 CPC — bold + UNDERLINED)
- Tables: tblLayout=fixed; first-row bold centered; explicit cell margins
- Page numbers: centered at TOP of page (Bombay HC convention per the
  filed pleading — pages numbered at top center, not bottom)

NOTE: Even with tblLayout=fixed locked, pandoc's pipe-table column widths are
not reliably honoured. The Drafter MUST run the post-pandoc `fix_docx_tables.py`
script after pandoc conversion to force the correct column widths.
"""
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("reference.docx")


def lock_style(style, *, font="Times New Roman", size_pt=14, bold=False,
               underline=False, align=None, line_spacing=1.5, color="000000",
               letter_spacing_pt=None, space_before_pt=0, space_after_pt=0,
               first_line_indent_cm=None, keep_with_next=False,
               outline_level=None):
    font_obj = style.font
    font_obj.name = font
    font_obj.size = Pt(size_pt)
    font_obj.bold = bold
    if underline:
        font_obj.underline = True
    if color:
        font_obj.color.rgb = RGBColor.from_string(color)

    # Force east-asian / hAnsi / cs font name locks
    rpr = style.element.get_or_add_rPr()
    existing = rpr.find(qn("w:rFonts"))
    if existing is not None:
        rpr.remove(existing)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)
    rFonts.set(qn("w:eastAsia"), font)
    rpr.append(rFonts)

    if letter_spacing_pt is not None:
        existing_sp = rpr.find(qn("w:spacing"))
        if existing_sp is not None:
            rpr.remove(existing_sp)
        spacing_el = OxmlElement("w:spacing")
        spacing_el.set(qn("w:val"), str(int(letter_spacing_pt * 20)))
        rpr.append(spacing_el)

    # Explicit underline (in case font.underline boolean doesn't stick)
    if underline:
        existing_u = rpr.find(qn("w:u"))
        if existing_u is not None:
            rpr.remove(existing_u)
        u_el = OxmlElement("w:u")
        u_el.set(qn("w:val"), "single")
        rpr.append(u_el)

    pf = style.paragraph_format
    if align is not None:
        pf.alignment = align
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    if first_line_indent_cm is not None:
        pf.first_line_indent = Cm(first_line_indent_cm)
    pf.keep_with_next = keep_with_next

    if outline_level is not None:
        ppr = style.element.get_or_add_pPr()
        existing_ol = ppr.find(qn("w:outlineLvl"))
        if existing_ol is not None:
            ppr.remove(existing_ol)
        ol = OxmlElement("w:outlineLvl")
        ol.set(qn("w:val"), str(outline_level))
        ppr.append(ol)


def main():
    doc = Document()

    # A4 + Bombay HC margins
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(4.0)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Normal (body)
    normal = doc.styles["Normal"]
    lock_style(normal, font="Times New Roman", size_pt=14, bold=False,
               align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5,
               first_line_indent_cm=0.5, space_after_pt=6)

    # Heading 1: Court header + case-number + cover-page anchors
    # Bold + centered, NOT underlined (matches "IN THIS HON'BLE COURT..." style)
    h1 = doc.styles["Heading 1"]
    lock_style(h1, font="Times New Roman", size_pt=14, bold=True,
               underline=False, align=WD_ALIGN_PARAGRAPH.CENTER,
               line_spacing=1.5, color="000000",
               space_before_pt=12, space_after_pt=12, keep_with_next=True,
               outline_level=0)
    h1.paragraph_format.first_line_indent = Cm(0)

    # Heading 2: spaced section headers — F A C T S, G R O U N D S, P R A Y E R,
    # I N D E X, S Y N O P S I S, L I S T   O F   A N N E X U R E S
    # Bold + UNDERLINED + centered + letter-spacing
    h2 = doc.styles["Heading 2"]
    lock_style(h2, font="Times New Roman", size_pt=14, bold=True,
               underline=True, align=WD_ALIGN_PARAGRAPH.CENTER,
               line_spacing=1.5, color="000000",
               letter_spacing_pt=4,
               space_before_pt=18, space_after_pt=6, keep_with_next=True,
               outline_level=1)
    h2.paragraph_format.first_line_indent = Cm(0)

    # Heading 3: unspaced section headers — SUBSTANTIAL QUESTIONS OF LAW,
    # ACTS & RULES, CITATIONS, MOST RESPECTFULLY SHEWETH, statutory openings
    # Bold + UNDERLINED + centered (Drafter can override alignment via inline style)
    h3 = doc.styles["Heading 3"]
    lock_style(h3, font="Times New Roman", size_pt=14, bold=True,
               underline=True, align=WD_ALIGN_PARAGRAPH.CENTER,
               line_spacing=1.5, color="000000",
               space_before_pt=12, space_after_pt=6, keep_with_next=True,
               outline_level=2)
    h3.paragraph_format.first_line_indent = Cm(0)

    # Heading 4: left-aligned bold underlined (for "MOST RESPECTFULLY SHEWETH:" type)
    h4 = doc.styles["Heading 4"]
    lock_style(h4, font="Times New Roman", size_pt=14, bold=True,
               underline=True, align=WD_ALIGN_PARAGRAPH.LEFT,
               line_spacing=1.5, color="000000",
               space_before_pt=12, space_after_pt=6, keep_with_next=True,
               outline_level=3)
    h4.paragraph_format.first_line_indent = Cm(0)

    # Title style — for statutory opening like "SECOND APPEAL UNDER SECTION 100..."
    title = doc.styles["Title"]
    lock_style(title, font="Times New Roman", size_pt=14, bold=True,
               underline=True, align=WD_ALIGN_PARAGRAPH.CENTER,
               line_spacing=1.5, color="000000",
               space_before_pt=12, space_after_pt=12, keep_with_next=True)
    title.paragraph_format.first_line_indent = Cm(0)

    # Body Text (sometimes used by pandoc after Heading 1)
    if "Body Text" in [s.name for s in doc.styles]:
        bt = doc.styles["Body Text"]
        lock_style(bt, font="Times New Roman", size_pt=14, bold=False,
                   align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5,
                   first_line_indent_cm=0.5, space_after_pt=6)

    # List Paragraph
    try:
        lp = doc.styles["List Paragraph"]
        lock_style(lp, font="Times New Roman", size_pt=14, bold=False,
                   align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5,
                   first_line_indent_cm=0, space_after_pt=6)
    except KeyError:
        pass

    # Page numbers at TOP CENTER (Bombay HC convention per the
    # filed pleading — pages numbered at top, not bottom)
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # Demo content for visual check + style anchor — pandoc replaces with draft
    doc.add_paragraph("IN THE HIGH COURT OF JUDICATURE AT BOMBAY BENCH AT [BENCH CITY].",
                      style="Heading 1")
    doc.add_paragraph("WRIT PETITION NO. _______ OF 2026", style="Heading 1")
    p = doc.add_paragraph("Style template — pandoc replaces this body. Normal style: TNR 14pt 1.5 spacing justified 0.5cm first-line indent.")
    doc.add_paragraph("F A C T S", style="Heading 2")
    doc.add_paragraph("Body inside F A C T S section.")
    doc.add_paragraph("ACTS & RULES", style="Heading 3")
    doc.add_paragraph("Indian Evidence Act 1872 / Civil Procedure Code 1908")
    doc.add_paragraph("MOST RESPECTFULLY SHEWETH:", style="Heading 4")
    doc.add_paragraph("Body following Heading 4.")

    doc.save(str(OUT_PATH))
    print(f"OK · wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
